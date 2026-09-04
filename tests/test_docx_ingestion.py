from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document

from app.knowledge.docx import DOCX_MIME_TYPE, DocxIngestor, DocxValidationError


def _document_bytes() -> bytes:
    document = Document()
    document.core_properties.title = "Guide ESS"
    document.add_heading("Reset password", level=1)
    document.add_paragraph("Open Employee Self Service and select Forgot Password.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Step"
    table.cell(0, 1).text = "Enter your username"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_ingestion_extracts_headings_tables_and_deterministic_chunks():
    ingestor = DocxIngestor()
    first = ingestor.ingest(
        _document_bytes(),
        filename="guide.docx",
        content_type=DOCX_MIME_TYPE,
        module_namespace="spay",
    )
    second = ingestor.ingest(
        _document_bytes(),
        filename="guide.docx",
        content_type=DOCX_MIME_TYPE,
        module_namespace="spay",
    )

    assert first.title == "Guide ESS"
    assert first.table_count == 1
    assert first.sections[0].heading == "Reset password"
    assert "Enter your username" in first.extracted_text
    assert first.chunks
    assert [chunk.chunk_id for chunk in first.chunks] == [
        chunk.chunk_id for chunk in second.chunks
    ]


def test_docx_rejects_non_word_zip():
    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w") as archive:
        archive.writestr("anything.txt", "not Word")

    with pytest.raises(DocxValidationError, match="Word document structure"):
        DocxIngestor().ingest(
            payload.getvalue(),
            filename="fake.docx",
            content_type=DOCX_MIME_TYPE,
            module_namespace="spay",
        )


def test_docx_rejects_macro_payload():
    data = _document_bytes()
    source = zipfile.ZipFile(io.BytesIO(data))
    payload = io.BytesIO()
    with source, zipfile.ZipFile(payload, "w") as output:
        for info in source.infolist():
            output.writestr(info, source.read(info.filename))
        output.writestr("word/vbaProject.bin", b"macro")

    with pytest.raises(DocxValidationError, match="Macro-enabled"):
        DocxIngestor().ingest(
            payload.getvalue(),
            filename="macro.docx",
            content_type=DOCX_MIME_TYPE,
            module_namespace="spay",
        )
